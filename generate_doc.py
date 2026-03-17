from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_toc(doc):
    paragraph = doc.add_paragraph()
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fld_simple)


doc = Document()

# Title
doc.add_heading('Network & Outreach Copilot – Milestone Plan', level=0)

# Table of Contents placeholder (Word built-in field)
add_toc(doc)

# Intro
p = doc.add_paragraph()
p.add_run('This document outlines the milestones to build the Network & Outreach Copilot from a basic prototype to an advanced product. It is structured so that each milestone is shippable and builds on the previous one.').bold = False

# Milestone 0
add_heading(doc, 'Milestone 0 – Foundations & Skeleton', level=1)

doc.add_paragraph('Goal: Set up the core stack and project skeleton.', style='List Bullet')

add_heading(doc, 'Stack Decisions', level=2)
stack = [
    'Frontend: Next.js (TypeScript, App Router) deployed on Vercel.',
    'Auth: Clerk for authentication (email/password, Google).',
    'Database/Backend: Supabase (Postgres, Row Level Security).',
    'Email: Resend for sending outreach and follow-up emails.',
    'Analytics: PostHog for product analytics.',
    'Error Tracking: Sentry for monitoring runtime errors.',
]
for item in stack:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_0 = [
    'Next.js project bootstrapped with TypeScript.',
    'Clerk authentication integrated with protected app routes.',
    'Supabase client configured and connection verified.',
    'Basic UI shell: layout, navigation, button, input, and card components.',
]
for item in deliverables_0:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 1
add_heading(doc, 'Milestone 1 – Basic Outreach Workspace (V1)', level=1)

doc.add_paragraph('Goal: Allow a student to set up a profile, add contacts, and generate a first personalized cold email.', style='List Bullet')

add_heading(doc, 'Core Features', level=2)
core_features_1 = [
    'Profile setup with name, school, degree, graduation year, interests, and experience level.',
    'Manual contact creation including name, role, company, email, LinkedIn URL, and notes.',
    'Opportunity definition with role type and free-text description or pasted job link.',
    'Outreach email generation (subject line and body) based on student profile, contact, and opportunity.',
    'Editable email text area with the ability to save generated messages.',
]
for item in core_features_1:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Initial Data Model', level=2)
data_model_1 = [
    'users: stores Clerk user id and profile fields.',
    'contacts: stores per-user contacts and metadata.',
    'opportunities: stores potential roles or situations tied to a user and optionally a contact.',
    'messages: stores generated outreach content with type, timestamps, and associations.',
]
for item in data_model_1:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_1 = [
    'Authenticated users can create and edit a personal profile.',
    'Users can add, view, and edit contacts.',
    'Users can define opportunities linked to contacts.',
    'Users can generate, edit, and save outreach emails.',
]
for item in deliverables_1:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 2
add_heading(doc, 'Milestone 2 – Email Sending & Basic Tracking (V1.5)', level=1)

doc.add_paragraph('Goal: Send emails directly from the app and track basic outreach history per contact.', style='List Bullet')

add_heading(doc, 'Features', level=2)
features_2 = [
    'Integrate Resend to send emails from within the application.',
    'Store send metadata such as sent_at timestamp and provider message identifiers.',
    'Add a simple contact CRM view showing last message date and a basic status tag.',
    'Per-contact timeline listing all messages and their timestamps.',
    'First follow-up generation based on the original outreach email and elapsed time.',
]
for item in features_2:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_2 = [
    'Send outreach emails via Resend from the message editor.',
    'View sent messages and basic status per contact.',
    'Generate and send at least one follow-up email per outreach.',
]
for item in deliverables_2:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 3
add_heading(doc, 'Milestone 3 – Role and Persona-Aware Messaging (V2)', level=1)

doc.add_paragraph('Goal: Make outreach clearly tailored to ML and data roles and to different contact personas.', style='List Bullet')

add_heading(doc, 'Persona Modeling', level=2)
persona_3 = [
    'Add persona tags to contacts such as industry machine learning engineer, researcher, data scientist, recruiter, and alumni.',
    'Capture opportunity attributes including seniority and domain focus (for example NLP, CV, ML infrastructure, analytics).',
]
for item in persona_3:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Messaging Improvements', level=2)
msg_improvements_3 = [
    'Use student profile, persona, opportunity details, and one or two projects as inputs to generation.',
    'Define internal templates for different outreach styles, such as short intro and ask, coffee chat request, and referral request.',
    'Allow users to select a template style before generation.',
]
for item in msg_improvements_3:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_3 = [
    'Contacts can be tagged with personas and opportunities with seniority and domain.',
    'Generated messages vary based on persona and template style.',
    'Users can save and reuse favorite templates.',
]
for item in deliverables_3:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 4
add_heading(doc, 'Milestone 4 – Projects and GitHub Integration (V3)', level=1)

doc.add_paragraph('Goal: Automatically include relevant projects and real work in outreach messages.', style='List Bullet')

add_heading(doc, 'Project Library', level=2)
project_library_4 = [
    'Create a projects table with title, description, tech stack, problem, links, and impact.',
    'Provide UI to add and edit projects and tag them by domain such as NLP, CV, tabular, recommender systems, and MLOps.',
]
for item in project_library_4:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'GitHub Integration', level=2)
github_4 = [
    'Allow users to connect GitHub or provide repository URLs.',
    'Fetch repository metadata such as name, description, stars, and last updated timestamp.',
    'Let users map repositories to projects and enrich them with additional context.',
]
for item in github_4:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Messaging Enhancements', level=2)
msg_4 = [
    'Automatically select one or two best matching projects for a given persona and role.',
    'Generate concise project references inside emails that highlight problem, approach, and impact.',
    'Avoid repetitive project mentions when contacting the same person multiple times.',
]
for item in msg_4:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_4 = [
    'Students maintain a structured project library.',
    'Messages reference relevant projects automatically based on role and persona.',
]
for item in deliverables_4:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 5
add_heading(doc, 'Milestone 5 – Follow-up Scheduling and Light Automation (V3.5)', level=1)

doc.add_paragraph('Goal: Help students remember to follow up and manage simple outreach sequences.', style='List Bullet')

add_heading(doc, 'Scheduling', level=2)
scheduling_5 = [
    'Define outreach sequences per contact and opportunity with steps such as outreach, follow-up one, and follow-up two.',
    'Use Upstash or scheduled jobs to check for due follow-ups.',
    'Notify users when a follow-up is due through email or in-app alerts.',
]
for item in scheduling_5:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Semi-automation and UI', level=2)
auto_5 = [
    'Pre-generate follow-up drafts when they become due and present them for review and send.',
    'Provide a simple pipeline view that shows the current stage of each opportunity.',
]
for item in auto_5:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_5 = [
    'Consistent reminders for follow-ups across contacts and opportunities.',
    'A clear per-opportunity pipeline visualizing outreach progress.',
]
for item in deliverables_5:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 6
add_heading(doc, 'Milestone 6 – Analytics and Learning What Works (V4)', level=1)

doc.add_paragraph('Goal: Use outcomes data to improve outreach quality and provide insights.', style='List Bullet')

add_heading(doc, 'Outcome Tracking', level=2)
outcomes_6 = [
    'Add outcome fields to opportunities and contacts such as no response, positive reply, neutral, negative, and interview.',
    'Start with manual outcome input by users.',
]
for item in outcomes_6:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Analytics and Insights', level=2)
analytics_6 = [
    'Integrate PostHog to track templates, styles, personas, and outcomes.',
    'Build a simple analytics page summarizing response rates by message style, length, and persona.',
    'Provide user-facing insights such as which styles and lengths are performing best.',
]
for item in analytics_6:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Adaptive Suggestions', level=2)
adaptive_6 = [
    'Before generating new outreach, surface suggestions based on what has historically worked for the user.',
    'Default templates and styles should gradually adapt to successful patterns.',
]
for item in adaptive_6:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_6 = [
    'Outcome tracking across contacts and opportunities.',
    'An analytics view that highlights effective outreach patterns.',
    'Generation defaults influenced by previous successful outreach.',
]
for item in deliverables_6:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 7
add_heading(doc, 'Milestone 7 – Intelligence Layer and Pinecone Search (V5)', level=1)

doc.add_paragraph('Goal: Make the copilot feel like a smart assistant that suggests who to contact and what to reference.', style='List Bullet')

add_heading(doc, 'Pinecone Integration', level=2)
pinecone_7 = [
    'Store vector embeddings for contacts, projects, and opportunities based on their textual descriptions.',
    'Use semantic search to suggest the best projects to reference for a given contact or role.',
    'Suggest similar contacts or roles that may be worth contacting in a small campaign.',
]
for item in pinecone_7:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Next-contact Suggestions and Campaigns', level=2)
next_contact_7 = [
    'Given a high-level target, rank existing contacts by relevance and prompt the user with recommended next steps.',
    'Allow users to create small, focused outreach campaigns with slightly varied yet consistent messages.',
]
for item in next_contact_7:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Deliverables', level=2)
deliverables_7 = [
    'Semantic search-powered suggestions for contacts and projects.',
    'Campaign tooling for targeted outreach.',
]
for item in deliverables_7:
    doc.add_paragraph(item, style='List Bullet')

# Milestone 8
add_heading(doc, 'Milestone 8 – Polishing and Growth Features', level=1)

doc.add_paragraph('Goal: Refine the product experience and add features that help it scale to more students and mentors.', style='List Bullet')

add_heading(doc, 'Enhancements', level=2)
enhancements_8 = [
    'Advisor and mentor review mode for suggesting edits to student messages.',
    'Preset playbooks such as alumni coffee chat templates and recruiter outreach templates for ML internships.',
    'A content library of high-performing messages that users can reuse and adapt.',
]
for item in enhancements_8:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, 'Overall Outcome', level=2)
doc.add_paragraph('By completing these milestones, the Network & Outreach Copilot evolves from a basic email generator into an intelligent assistant that helps ML and data students systematically build and leverage their professional network.', style='List Bullet')


doc.save('documentation_v1.docx')

